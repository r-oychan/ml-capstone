from docx import Document
from docx.shared import Pt

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

doc.add_heading('BBO Capstone Reflection', level=0)

doc.add_heading('1. Initial Codebase', level=1)
doc.add_paragraph(
    "I built my codebase from scratch. After some initial research into Bayesian "
    "Optimization and Gaussian Processes, it became clear that implementing the "
    "pipeline and model by hand was not overly difficult, and doing so gave me a "
    "much deeper understanding of how the pieces fit together — the GP surrogate, "
    "the Expected Improvement acquisition function, and the loop that proposes the "
    "next query point. Building it myself also meant I could shape the code around "
    "my workflow rather than bending my workflow to fit someone else's library. "
    "My repository is publicly available at: https://github.com/r-oychan/ml-capstone"
)

doc.add_heading('2. Code Modification', level=1)
doc.add_paragraph(
    "The code evolved iteratively across rounds. Early on, the focus was on the "
    "core BO logic — GP fitting, EI, and a simple proposal function. As rounds "
    "went on, I added supporting infrastructure rather than rewriting the model:"
)
doc.add_paragraph(
    "• Folder structure and helper functions to hold each round's submission "
    "inputs and outputs (txt files), and to load them cleanly back into the "
    "main program. This made the pipeline far easier to maintain and reproduce.",
    style='List Bullet'
)
doc.add_paragraph(
    "• Logging tables summarising per-function best values and round-by-round "
    "history, so I could see at a glance which functions were improving and "
    "which were stuck.",
    style='List Bullet'
)
doc.add_paragraph(
    "• Graphing and visualisations of convergence and progress, which made "
    "it obvious when a function had plateaued and when exploitation was still "
    "paying off.",
    style='List Bullet'
)
doc.add_paragraph(
    "These changes didn't alter the math of the BO loop, but they had a big "
    "impact on my decision-making: they let me adjust strategy and the xi "
    "exploration parameter each round based on real evidence rather than guesswork. "
    "The single most significant change was moving from a uniform, one-size-fits-all "
    "strategy to per-function tuning in Round 4 — different xi, radius, and "
    "local/broad candidate mixes for each function. That round produced three new "
    "best values and validated that the functions genuinely needed different "
    "treatments."
)

doc.add_heading('3. Final Result', level=1)
doc.add_paragraph(
    "Most functions progressed over time, but progress was not monotonic — even "
    "after a new best, further exploitation of that region did not always produce "
    "additional gains. With only 13 rounds total, balancing exploration and "
    "exploitation was genuinely hard: every round spent exploring was a round not "
    "spent refining a known good region, and vice versa."
)
doc.add_paragraph(
    "In hindsight, if I had more time or a fresh start, I would:"
)
doc.add_paragraph(
    "• Lean harder into exploitation on functions showing a clear upward trend, "
    "rather than hedging with broad exploration.",
    style='List Bullet'
)
doc.add_paragraph(
    "• Use more systematic exploration — closer to a grid or low-discrepancy "
    "sequence — on functions that never moved (F1 in particular). Random broad "
    "sampling never surfaced a second peak for F1, and a structured sweep might "
    "have.",
    style='List Bullet'
)
doc.add_paragraph(
    "• Commit to a strategy earlier and resist the temptation to re-tune every "
    "round; discipline probably beats constant adjustment with such a small "
    "budget.",
    style='List Bullet'
)

doc.add_heading('4. Trade-offs and Decisions', level=1)
doc.add_paragraph(
    "The biggest trade-off was exploration versus exploitation under a severe "
    "budget constraint. F1 is the clearest example: it never really moved, and "
    "spending repeated rounds exploiting near its initial best felt wasteful — "
    "but of course I could not know that in advance. In general, with only 13 "
    "rounds, weighting more heavily toward exploitation felt like the safer bet: "
    "you can't discover much with a handful of exploration queries in higher "
    "dimensions anyway."
)
doc.add_paragraph(
    "If the budget had been larger, I would have shifted the weighting toward "
    "more exploration early on. A connected lesson is that staying disciplined "
    "with a consistent weighting — rather than second-guessing every round — is "
    "probably a better long-run approach. The short-term temptation is always to "
    "react to the last result; the long-term strategy is to trust the framework "
    "and let the GP accumulate information."
)

doc.add_heading('5. Learning and Application', level=1)

doc.add_heading('Part A — Most Important Lesson', level=2)
doc.add_paragraph(
    "The most valuable part was simply building the whole thing from scratch — "
    "doing the coding end to end, from GP fitting to acquisition optimisation "
    "to the submission pipeline. On top of that, there was something genuinely "
    "fun about the loop itself: submitting, waiting, checking email to see "
    "whether any function had improved, and then adjusting strategy for the "
    "next round. It made the abstract idea of \"active learning under a budget\" "
    "concrete."
)
doc.add_paragraph(
    "The conceptual lesson — how to budget exploration versus exploitation "
    "under a fixed query budget — generalises widely. In reinforcement "
    "learning, it is the same fundamental tension. In my own domain, the "
    "payments industry, we traditionally train a supervised model and set a "
    "threshold to approve or decline transactions. An RL-style framing instead "
    "defines a reward function and deliberately injects exploration so the "
    "system can discover fraud patterns or approval behaviours that a purely "
    "exploitative policy would miss. The BBO project made that trade-off "
    "tangible in a way that reading about it did not."
)

doc.add_heading('Part B — What Surprised Me Most', level=2)
doc.add_paragraph(
    "The biggest surprise was how flat F1 turned out to be — the numbers barely "
    "moved across 13 attempts, and no amount of local precision or broad "
    "exploration shook anything loose. It was a reminder that with a tiny "
    "query budget, some functions simply won't yield, and you have to accept that."
)
doc.add_paragraph(
    "The other surprise came from seeing how differently my peers approached "
    "the same problem. Some used one notebook per function; I went with a "
    "single Python module driving everything. I used a ConstantKernel × "
    "Matern(2.5) + WhiteKernel combination — the ConstantKernel handling the "
    "amplitude scale and the WhiteKernel absorbing observation noise — which "
    "worked well on some functions and less so on others. What struck me most "
    "was seeing someone get strong results with a plain random forest instead "
    "of a GP. It was a useful reality check: the sophisticated tool isn't "
    "always the winning tool, and with this few data points, there's a real "
    "element of luck in which model happens to fit the hidden landscape best."
)

doc.save('/home/user/ml-capstone/bbo_capstone_reflection.docx')
print('Saved reflection docx')
